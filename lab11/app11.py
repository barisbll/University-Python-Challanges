from math import floor
import matplotlib.pyplot as plt
import docx, re
import urllib.request
from PIL import Image
from docx.shared import Pt, Inches


def run():

    book_name = 'donquixote.docx'

    book = read_book(book_name)

    title, author, chapter = search_book(book)
    #print(title)

    words = count_paragraphs(chapter)

    plot_words(words)

    download_image()

    resize_image()

    turn_and_paste_image()

    make_word_doc(title, author, chapter, words)


def read_book(book_name):
    doc = docx.Document(book_name)
    return doc


def search_book(book):

    title_regex = re.compile(r'Title: .*')
    title = ''

    author_regex = re.compile(r'Author: .*')
    author = ''

    chapter_found = False
    chapter1_regex = re.compile(r'^CHAPTER I.$')
    chapter2_regex = re.compile(r'CHAPTER II.')
    chapter_paragraphs = []

    for e in book.paragraphs:
        if title_regex.match(e.text):
            title = title_regex.search((e.text)).group()[7:]

        if author_regex.match(e.text):
            author = author_regex.search((e.text)).group()[8:]

        if chapter1_regex.match(e.text):
            chapter_found = True
            chapter_paragraphs.append(e)

        if chapter2_regex.match(e.text):
            chapter_found = False

        if chapter_found:
            chapter_paragraphs.append(e)


    return (title, author, chapter_paragraphs)


def count_paragraphs(chapter):

    words = []

    for e in chapter:
        words.append(floor(len(e.text.split(' '))/10))

    return words


def plot_words(words):

    plt.scatter(x=range(0, len(words)), y=words)
    plt.savefig('plot.png')


def download_image():
    urllib.request.urlretrieve('https://cdn.oggito.com/images/full/2019/11/quijotedon2.jpg',
                               'don1.jpg')


def resize_image():
    img = Image.open('don1.jpg')

    # Size of the image in pixels (size of orginal image)
    # (This is not mandatory)
    width, height = img.size

    # Setting the points for cropped image
    left = width/3
    top = height / 8
    right = width - width/5
    bottom = 2 * height / 4

    # Cropped image of above dimension
    # (It will not change orginal image)
    img = img.crop((left, top, right, bottom))

    img = img.resize((width*2, height*2))


    img.save('don1_modified.jpg')


def turn_and_paste_image():
    img2 = Image.open('don2.png')
    img2 = img2.rotate(-30, expand=True)
    w2, h2 = img2.size
    img2 = img2.resize((int(w2/3), int(h2/3)))

    img1 = Image.open('don1_modified.jpg')


    img1.paste(img2, (0,570))

    img1.save('pasted_image.jpg')

# num paragraphs, num words 1.st chapter, min and max number of words in paragraphs, average number of words in paragraphs ,
# source of work
def make_word_doc(title, author, chapter, words):

    doc = docx.Document()
    style_norm = doc.styles['Normal']
    font_norm = style_norm.font
    font_norm.name = 'Calibri'
    font_norm.bold = True
    font_norm.size = Pt(16)

    heading = doc.add_heading('Report for lab:11', 0)
    style_header = heading.style
    font_header = style_header.font
    font_header.italic = True
    font_header.name = 'Arial'
    font_header.size = Pt(48)


    doc.add_paragraph(title,'Title')
    doc.add_paragraph(author)
    doc.add_paragraph('Baris Balli')

    doc.add_picture('plot.png')
    doc.add_picture('pasted_image.jpg', width=Inches(4.0))

    number_of_paragraphs = len(chapter)
    number_of_words = sum(words)

    count_zero_words = 0
    count_one_words = 0

    for e in words:
        if e == 0:
            count_zero_words += 1
        if e == 1:
            count_one_words += 1

    #words_real is used to find number of words per paragraph without yaking the floor

    words_real = []

    for e in chapter:
        words_real.append(len(e.text.split(' ')))

    min_number_of_words = min(words_real)
    max_number_of_words = max(words_real)
    avg_number_of_words = round(sum(words_real)/len(words_real),2)

    doc.add_paragraph(f'Number of paragraphs: {number_of_paragraphs}')
    doc.add_paragraph(f'Number of words in first chapter: {number_of_words}')
    doc.add_paragraph(f'Min number of words in first chapter: {min_number_of_words}')
    doc.add_paragraph(f'Max number of words in first chapter: {max_number_of_words}')
    doc.add_paragraph(f'Number of zero values in the plot: {count_zero_words}')
    doc.add_paragraph(f'Number of one values in the plot: {count_one_words}')
    doc.add_paragraph(f'Average number of words in first chapter: {avg_number_of_words}')
    doc.add_paragraph(f'Source for the book analyzed : https://www.gutenberg.org/ebooks/996')



    doc.save('report.docx')


run()
